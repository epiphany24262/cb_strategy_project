from __future__ import annotations

import base64
import re
from pathlib import Path

import nbformat
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
NOTEBOOK_PATH = ROOT / "notebooks" / "convertible_bond_factor_report.ipynb"
REPORT_MD_PATH = ROOT / "report" / "report.md"
ASSETS_DIR = ROOT / "report" / "assets"
DOCX_PATH = ROOT / "report" / "report.docx"
DOCX_FALLBACK_PATH = ROOT / "report" / "report_v2.docx"

FIGURE_SPECS = [
    {"cell": 5, "image_idx": 0, "filename": "fig01_universe_size.png"},
    {"cell": 15, "image_idx": 0, "filename": "fig02_dblow_group_nav.png"},
    {"cell": 15, "image_idx": 1, "filename": "fig03_bond_prem_group_nav.png"},
    {"cell": 15, "image_idx": 2, "filename": "fig04_alpha_pct_chg_5_group_nav.png"},
    {"cell": 21, "image_idx": 0, "filename": "fig05_factor_corr.png"},
    {"cell": 27, "image_idx": 0, "filename": "fig06_strategy_vs_benchmark.png"},
    {"cell": 30, "image_idx": 0, "filename": "fig07_excess_nav.png"},
    {"cell": 35, "image_idx": 1, "filename": "fig08_execution_compare.png"},
    {"cell": 30, "image_idx": 1, "filename": "fig09_hedge20_nav.png"},
    {"cell": 15, "image_idx": 3, "filename": "figA1_turnover_5_group_nav.png"},
    {"cell": 39, "image_idx": 0, "filename": "figA2_rolling_validation_nav.png"},
]

BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
ACCENT_COLOR = RGBColor(0x00, 0x00, 0x00)
MUTED_COLOR = RGBColor(0x44, 0x44, 0x44)
HEADER_FILL = "EDEDED"
ROW_FILL = "F7F7F7"
INFO_LEFT_FILL = "EEEEEE"
INFO_RIGHT_FILL = "FAFAFA"


def extract_notebook_images() -> dict[str, Path]:
    """Export selected notebook output images to report/assets."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    nb = nbformat.read(NOTEBOOK_PATH, as_version=4)
    saved: dict[str, Path] = {}

    for spec in FIGURE_SPECS:
        cell = nb.cells[spec["cell"]]
        image_count = 0
        image_data = None
        for output in cell.get("outputs", []):
            data = output.get("data", {})
            if "image/png" in data:
                if image_count == spec["image_idx"]:
                    image_data = data["image/png"]
                    break
                image_count += 1

        if image_data is None:
            raise RuntimeError(
                f"Cannot find image {spec['image_idx']} in notebook cell {spec['cell']}"
            )

        img_path = ASSETS_DIR / spec["filename"]
        img_path.write_bytes(base64.b64decode(image_data))
        saved[spec["filename"]] = img_path

    return saved


def sanitize_inline(text: str) -> str:
    """Remove lightweight Markdown markers for table cells and plain fallbacks."""
    return text.replace("**", "").replace("`", "")


def set_run_font(
    run,
    size_pt: float = 12,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    latin_font: str = "Times New Roman",
    east_asia_font: str = "宋体",
):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    run.font.color.rgb = color or BODY_COLOR

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def add_inline_runs(
    paragraph,
    text: str,
    *,
    size_pt: float = 12,
    base_bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    east_asia_font: str = "宋体",
):
    """Render **bold** and `code` markers while keeping the Word export simple."""
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()].replace("`", ""))
            set_run_font(
                run,
                size_pt=size_pt,
                bold=base_bold,
                italic=italic,
                color=color,
                east_asia_font=east_asia_font,
            )
        run = paragraph.add_run(match.group(1).replace("`", ""))
        set_run_font(
            run,
            size_pt=size_pt,
            bold=True,
            italic=italic,
            color=color,
            east_asia_font=east_asia_font,
        )
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:].replace("`", ""))
        set_run_font(
            run,
            size_pt=size_pt,
            bold=base_bold,
            italic=italic,
            color=color,
            east_asia_font=east_asia_font,
        )


def set_paragraph_bottom_border(paragraph, color: str = "808080", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)

    head = paragraph.add_run("第 ")
    set_run_font(head, size_pt=9.5, color=MUTED_COLOR)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    set_run_font(r, size_pt=9.5, color=MUTED_COLOR)
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)

    tail = paragraph.add_run(" 页")
    set_run_font(tail, size_pt=9.5, color=MUTED_COLOR)


def add_header(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("可转债多因子选券研究")
    set_run_font(run, size_pt=9.5, color=MUTED_COLOR)


def configure_document(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BODY_COLOR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    add_header(sec.header.paragraphs[0])
    add_page_number(sec.footer.paragraphs[0])


def add_horizontal_rule(doc: Document, color: str = "000000", size: str = "10"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_bottom_border(p, color=color, size=size)


def add_cover_page(doc: Document):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("可转债多因子选券研究报告")
    set_run_font(run, size_pt=24, bold=True, color=ACCENT_COLOR, east_asia_font="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("金融程序化交易")
    set_run_font(run, size_pt=14, color=MUTED_COLOR)

    add_horizontal_rule(doc)

    items = [
        ("研究对象", "A 股可转债横截面多因子选券"),
        ("数据区间", "2018-01-02 至 2024-12-26"),
        ("策略口径", "三因子评分 + amount q20 + BW 缓冲排名加权"),
        ("评估维度", "收益、稳定性、可交易性、可复现性"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for left, right in items:
        row = table.add_row().cells
        set_cell_shading(row[0], INFO_LEFT_FILL)
        set_cell_shading(row[1], INFO_RIGHT_FILL)
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(left)
        set_run_font(r0, size_pt=11, bold=True, color=ACCENT_COLOR, east_asia_font="黑体")
        p1 = row[1].paragraphs[0]
        r1 = p1.add_run(right)
        set_run_font(r1, size_pt=11)

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("课程作业最终版")
    set_run_font(run, size_pt=12, color=MUTED_COLOR)

    doc.add_page_break()


def add_quick_navigation_page(doc: Document):
    """Add a formal static table of contents page.

    Page numbers are fixed to the current report layout. If the report body
    changes substantially, update the page numbers in TOC_ENTRIES accordingly.
    """
    toc_entries = [
        (1, "0.1 主要结论", "3"),
        (1, "1. 数据与样本过滤", "3"),
        (2, "1.1 样本背景", "3"),
        (2, "1.2 样本过滤与可交易性约束", "4"),
        (1, "2. 单因子分组与 IC/IR", "5"),
        (2, "2.1 检验口径", "5"),
        (2, "2.2 dblow：双低因子", "5"),
        (2, "2.3 bond_prem：债底安全边际补充", "5"),
        (2, "2.4 alpha_pct_chg_5：正股短期联动修正", "6"),
        (2, "2.5 因子统计汇总", "7"),
        (1, "3. 多因子合成与相关性", "8"),
        (2, "3.1 合成方法", "8"),
        (2, "3.2 组合敏感性比较", "8"),
        (1, "4. 缓冲排名加权 BW 主策略", "10"),
        (2, "4.1 策略结构", "10"),
        (2, "4.2 绩效指标", "10"),
        (2, "4.3 多空 20% 对冲验证", "11"),
        (1, "5. 稳健性分析", "13"),
        (2, "5.1 调仓频率敏感性", "13"),
        (2, "5.2 TopN 敏感性", "13"),
        (2, "5.3 手续费敏感性", "13"),
        (2, "5.4 年度表现", "14"),
        (2, "5.5 分阶段表现", "14"),
        (2, "5.6 执行层对比", "14"),
        (2, "5.7 可比口径参数网格", "15"),
        (2, "5.8 滚动验证与执行阈值附录", "16"),
        (1, "6. 风险与局限性", "18"),
        (1, "7. 结论", "19"),
        (1, "附录", "19"),
        (2, "A.1 候选因子去留总表", "19"),
        (2, "A.2 turnover_5 单因子补充图", "20"),
        (2, "A.3 补充实验说明", "20"),
    ]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("目录")
    set_run_font(run, size_pt=18, bold=True, color=ACCENT_COLOR, east_asia_font="黑体")
    set_paragraph_bottom_border(p, color="000000", size="8")

    for level, title, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.65)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(14.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = p.add_run(f"{title}\t{page}")
        set_run_font(
            run,
            size_pt=11.2 if level == 1 else 10.5,
            bold=(level == 1),
            color=ACCENT_COLOR if level == 1 else BODY_COLOR,
            east_asia_font="黑体" if level == 1 else "宋体",
        )

    doc.add_page_break()

def add_text_paragraph(doc: Document, text: str, center: bool = False, size_pt: float = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if not center:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(p, text, size_pt=size_pt)
    return p


def add_lead_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text, size_pt=11.5, italic=True, color=MUTED_COLOR)


def add_heading(doc: Document, text: str, level: int):
    if level == 2 and re.match(r"^\d+\.\s", text) and not text.startswith("1. "):
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    p = doc.add_paragraph()

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(16)
        add_inline_runs(
            p,
            text,
            size_pt=20,
            base_bold=True,
            color=ACCENT_COLOR,
            east_asia_font="黑体",
        )
        return

    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6 if level in (2, 3) else 3)

    if level == 2:
        add_inline_runs(
            p,
            text,
            size_pt=14.5,
            base_bold=True,
            color=ACCENT_COLOR,
            east_asia_font="黑体",
        )
        set_paragraph_bottom_border(p)
    elif level == 3:
        add_inline_runs(
            p,
            text,
            size_pt=12.5,
            base_bold=True,
            color=ACCENT_COLOR,
            east_asia_font="黑体",
        )
    else:
        add_inline_runs(
            p,
            text,
            size_pt=11.5,
            base_bold=True,
            color=MUTED_COLOR,
            east_asia_font="黑体",
        )


def add_bullet(doc: Document, text: str, numbered: bool = False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, text, size_pt=12)


def add_image(doc: Document, img_path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(14.8))


def add_caption(doc: Document, text: str, italic: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2 if text.startswith(("图", "附图")) else 6)
    add_inline_runs(
        p,
        text,
        size_pt=10.5,
        base_bold=text.startswith(("图", "附图")),
        italic=italic,
        color=MUTED_COLOR,
    )


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    current = lines[idx].strip()
    nxt = lines[idx + 1].strip()
    return current.startswith("|") and nxt.startswith("|") and set(nxt.replace("|", "").strip()) <= {"-", ":", " "}


def parse_markdown_table(lines: list[str], idx: int):
    header = [x.strip() for x in lines[idx].strip().strip("|").split("|")]
    i = idx + 2
    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        rows.append([x.strip() for x in line.strip("|").split("|")])
        i += 1
    return header, rows, i


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    is_info_table = [sanitize_inline(x) for x in header] == ["项目", "内容"]

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(
            p,
            text,
            size_pt=11,
            base_bold=True,
            color=ACCENT_COLOR,
            east_asia_font="黑体",
        )

    for row_idx, row in enumerate(rows):
        tr = table.add_row().cells
        for j, text in enumerate(row):
            tr[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_info_table:
                set_cell_shading(tr[j], INFO_LEFT_FILL if j == 0 else INFO_RIGHT_FILL)
            elif row_idx % 2 == 1:
                set_cell_shading(tr[j], ROW_FILL)
            p = tr[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, text, size_pt=10.5)

    doc.add_paragraph()


def build_docx():
    assets = extract_notebook_images()
    text = REPORT_MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    add_cover_page(doc)
    add_quick_navigation_page(doc)

    paragraph_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_text_paragraph(doc, " ".join(x.strip() for x in paragraph_buffer if x.strip()))
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_lead_line(doc, stripped[2:].strip())
            i += 1
            continue

        if is_table_block(lines, i):
            flush_paragraph()
            header, rows, next_idx = parse_markdown_table(lines, i)
            add_table(doc, header, rows)
            i = next_idx
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            rel_path = image_match.group(2).strip()
            img_name = Path(rel_path).name
            img_path = assets.get(img_name, ROOT / "report" / rel_path)
            add_image(doc, img_path)
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            while i < len(lines) and lines[i].strip().startswith(("图", "附图", "注：")):
                add_caption(doc, lines[i].strip(), italic=False)
                i += 1
                while i < len(lines) and not lines[i].strip():
                    i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_bullet(doc, stripped[2:].strip(), numbered=False)
            i += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if number_match:
            flush_paragraph()
            add_bullet(doc, number_match.group(1).strip(), numbered=True)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        doc.save(DOCX_FALLBACK_PATH)
        return DOCX_FALLBACK_PATH


if __name__ == "__main__":
    output_path = build_docx()
    print(f"Generated {output_path}")
